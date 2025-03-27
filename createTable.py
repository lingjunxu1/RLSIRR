from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from functools import partial
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def createTable(dataNames):
    document = Document()
    document.styles['Normal'].font.name = u'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')
    #p=document.add_paragraph()
    #p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    table = document.add_table(rows=32,cols=5,style='Table Grid')#生成rows*cols的表格
    for i in range(32):
        for j in range(5):
            Bold(table.rows[i].cells[j])
            clear_LR(table.rows[i].cells[j])
    #table.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    table.autofit = False
    table.cell(0,0).merge(table.cell(1,0))
    table.cell(0,1).merge(table.cell(1,1))
    table.cell(0,2).merge(table.cell(0,4))
    for i in range(2,27+1,5):
        table.cell(i,0).merge(table.cell(i+4,0))
    data = ["DataSet"]+["\n\n"+i for i in dataNames]
    index = ["SSIM      ↑","PSNR     ↑","NCC       ↑","LMSE    ↓","Distance ↓"]
    model = ["Input","Original","Current"]
    table.rows[0].cells[0].text = data[0]
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    #table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True   ★ 文本加粗就用这行
    #table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(250,0,0) #★ 换颜色就用这句
    for i in range(6):
        table.rows[2+i*5].cells[0].text = data[i+1]
        table.rows[2+i*5].cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    table.rows[0].cells[1].text = "Index"
    table.rows[0].cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i in range(6):
        for j in range(5):
            table.rows[2+i*5+j].cells[1].text = index[j]
            table.rows[2+i*5+j].cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.rows[0].cells[2].text = "Model"
    table.rows[0].cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for i,j in enumerate([2,3,4]): 
        table.rows[1].cells[j].text = model[i]
        table.rows[1].cells[j].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return document,table



def set_cell_border(cell: _Cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV','left','right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
clear_LR = partial(set_cell_border, left={"sz":1 , "val": "dashed", "color": "FFFFFF","shadow": "true"},right={"sz": 1, "color": "FFFFFF","val": "dashed"})
clear_B = partial(set_cell_border,bottom={"sz":12, "color": "FFFFFF", "val": "single"})
Bold = partial(set_cell_border,top={"sz":10,"color": "000000","val": "single"},bottom={"sz":10,"color": "000000","val": "single"},left={"sz":10,"color": "000000","val": "single"},right={"sz":10,"color": "000000","val": "single"})
color = [[(255,0,0),(50,255,0)],[(255,0,0),(50,255,0)],[(255,0,0),(50,255,0)],[(50,255,0),(255,0,0)],[(50,255,0),(255,0,0)]]*6

def saveFile(dataSetNames,dataNum,testData,savePath,flagPath,saveName = "results"):
    document,table = createTable([key+f"({dataNum[i]})" for i,key in enumerate(dataSetNames)]+[f"Average({sum(dataNum)})"])
    indexName = ["ssim","psnr","ncc","lmse","distance"]
    modelName = ["Input","Original","Current"]
    count = 0
    averange = {model:{i:0 for i in indexName} for model in modelName}
    data = [[0 for i in range(len(modelName))] for j in range(5*(len(dataSetNames)+1))]
    for datasetName in testData:
        index = testData[datasetName]
        for i,model in enumerate(modelName):
            for j,name in enumerate(indexName):
                num = index[model][name]
                data[count*5+j][i] = num
                averange[model][name]+=num*dataNum[count]
        count+=1
    for i,model in enumerate(modelName):
        for j,name in enumerate(indexName):
            num = averange[model][name]/sum(dataNum)
            data[count*5+j][i] = num
    flagFlag = sum([[0+5*ii,1+5*ii,2+5*ii] for ii in range(6)],[])
    for i,line in enumerate(data):
        temp = [round(num,3) if num<=1 else round(num,2) for num in line]
        if i in flagFlag:Num = max(temp)
        else: Num = min(temp)
        flag = [num==Num for num in temp]
        for j,num in enumerate(line):
            if i not in [4,9,14,19,24,29]:
                clear_B(table.rows[i+2].cells[2+j])
                if j==0: clear_B(table.rows[i+2].cells[1+j])
            #if num<=1: table.rows[i+2].cells[2+j].text = str(round(num,3))
            #else:      table.rows[i+2].cells[2+j].text = str(round(num,2))
            if num<=1:run = table.rows[i+2].cells[j+2].paragraphs[0].add_run(str(round(num,3)))
            else:     run = table.rows[i+2].cells[j+2].paragraphs[0].add_run(str(round(num,2)))
            if flag[j]: 
                run.font.bold = True
                run.underline = True
            table.rows[i+2].cells[j+2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if j in [2]:
                if data[i][j]>=1: dk = 2
                else: dk=3
                differ = round(data[i][j],dk) - round(data[i][j-1],dk)
                #if data[i][j]>=1: dk = 2
                #else: dk=3
                #print(differ)
                if differ==0:#round(differ,dk)==0: 
                    rgb = RGBColor(0,50,255)
                    s = "   -------"
                elif differ>0: 
                    rgb = RGBColor(*color[i][0])
                    s = " ↑{:.3f}".format(abs(differ))
                else: 
                    rgb = RGBColor(*color[i][1])
                    s = " ↓{:.3f}".format(abs(differ))
                run = table.rows[i+2].cells[j+2].paragraphs[0].add_run(s)
                run.font.color.rgb = rgb
    if not os.path.exists(savePath): os.makedirs(savePath)
    document.add_paragraph("")
    document.add_paragraph(flagPath)
    
    document.save(os.path.join(savePath,f"{saveName}.docx"))
    os.system("libreoffice --invisible --convert-to pdf --outdir {} {}.docx".format(savePath,os.path.join(savePath,saveName)))
    os.remove(os.path.join(savePath,f"{saveName}.docx"))
